from paddleocr import PaddleOCR
import cv2
from docx import Document
from docx.shared import Pt, RGBColor
from pathlib import Path
import os
import re
import numpy as np
import time

def light_preprocess(img_path):
    img = cv2.imread(img_path)
    if img is None:
        return None
    
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    clahe = cv2.createCLAHE(clipLimit=1.5, tileGridSize=(8,8))
    enhanced = clahe.apply(gray)
    denoised = cv2.fastNlMeansDenoising(enhanced, h=7)
    
    return denoised

def split_long_image(img_path, chunk_height=3000):
    img = cv2.imread(str(img_path))
    
    if img is None:
        return []
    
    h, w = img.shape[:2]
    
    chunks = []
    overlap = 200
    
    for y in range(0, h, chunk_height - overlap):
        y_end = min(y + chunk_height, h)
        chunk = img[y:y_end, :]
        
        temp_path = f'temp_chunk_{len(chunks)}.jpg'
        cv2.imwrite(temp_path, chunk)
        
        preprocessed = light_preprocess(temp_path)
        if preprocessed is not None:
            temp_preprocessed = f'temp_prep_{len(chunks)}.jpg'
            cv2.imwrite(temp_preprocessed, preprocessed)
            chunks.append((temp_preprocessed, temp_path))
        else:
            chunks.append((temp_path, temp_path))
    
    return chunks

def smart_postprocess(text):
    text = re.sub(r'\s{2,}', ' ', text)
    return text.strip()

def is_noise(text):
    text = text.strip()
    if len(text) < 2:
        return True
    
    if re.match(r'^[0-9<>|\.\-_~*:\'"]+$', text):
        return True
    
    noise_patterns = ['UU', 'Mdo', '00000', '|0', '::', 'H']
    if text in noise_patterns:
        return True
    
    return False

def group_texts_by_proximity(texts_with_boxes, vertical_threshold=100):
    if not texts_with_boxes:
        return []
    
    sorted_texts = sorted(texts_with_boxes, key=lambda x: x[1])
    bubbles = []
    current_bubble = [sorted_texts[0]]
    
    for i in range(1, len(sorted_texts)):
        prev_y = current_bubble[-1][1]
        curr_y = sorted_texts[i][1]
        
        if abs(curr_y - prev_y) < vertical_threshold:
            current_bubble.append(sorted_texts[i])
        else:
            bubbles.append(current_bubble)
            current_bubble = [sorted_texts[i]]
    
    if current_bubble:
        bubbles.append(current_bubble)
    
    result = []
    for bubble in bubbles:
        combined_text = ' '.join([text for text, y, score in bubble])
        combined_text = smart_postprocess(combined_text)
        avg_score = sum([score for text, y, score in bubble]) / len(bubble)
        result.append((combined_text, avg_score))
    
    return result

print('='*70)
print('EXTRACTION COMPLETE - CHAPITRE 35'.center(70))
print('='*70)

start_time = time.time()

print('\nInitialisation de PaddleOCR...')
ocr_ko = PaddleOCR(
    lang='korean',
    use_textline_orientation=True,
    text_det_thresh=0.25,
    text_det_box_thresh=0.55,
    text_det_unclip_ratio=1.6
)

# Dossier source
folder = Path(r'P:\19 - The Detective Agency For Regretful Male Leads (1)\Chapitre 35')

# Trouver toutes les images
images = sorted(folder.glob('*.jpeg')) + sorted(folder.glob('*.jpg')) + sorted(folder.glob('*.png'))
images = sorted(set(images))

print(f'\n{len(images)} pages trouvees\n')
print('='*70)

# Document Word
doc = Document()
doc.add_heading('The Detective Agency', 0)
doc.add_heading('Chapitre 35 - Version Coréenne', 1)
doc.add_paragraph(f'Extraction de {len(images)} pages')
doc.add_paragraph()

total_bubbles = 0

for page_num, img_path in enumerate(images, 1):
    print(f'\n[PAGE {page_num}/{len(images)}] {img_path.name}')
    
    # Lire dimensions
    img = cv2.imread(str(img_path))
    if img is None:
        print('  ERREUR: Impossible de lire l\'image')
        continue
    
    h, w = img.shape[:2]
    print(f'  Taille: {w}x{h} pixels')
    
    # Titre dans le document
    doc.add_heading(f'Page {page_num}', level=2)
    
    # Decouper si necessaire
    chunks = split_long_image(img_path)
    if len(chunks) > 1:
        print(f'  Decoupee en {len(chunks)} morceaux')
    
    page_bubbles = []
    
    for prep_path, orig_path in chunks:
        try:
            result = ocr_ko.predict(prep_path)
            
            if result and len(result) > 0:
                for page_result in result:
                    if 'rec_texts' in page_result and 'rec_boxes' in page_result:
                        texts = page_result['rec_texts']
                        scores = page_result['rec_scores']
                        boxes = page_result['rec_boxes']
                        
                        texts_with_positions = []
                        for text, score, box in zip(texts, scores, boxes):
                            if score > 0.70 and text.strip() and not is_noise(text):
                                y_pos = box[1]
                                texts_with_positions.append((text, y_pos, score))
                        
                        bubbles = group_texts_by_proximity(texts_with_positions)
                        
                        for bubble_text, avg_score in bubbles:
                            page_bubbles.append(bubble_text)
            
            # Nettoyer
            if os.path.exists(prep_path):
                os.remove(prep_path)
            if os.path.exists(orig_path) and orig_path != prep_path:
                os.remove(orig_path)
                
        except Exception as e:
            print(f'  Erreur: {e}')
    
    print(f'  -> {len(page_bubbles)} bulles detectees')
    total_bubbles += len(page_bubbles)
    
    # Ajouter au document
    if page_bubbles:
        for idx, text in enumerate(page_bubbles, 1):
            p = doc.add_paragraph()
            
            num_run = p.add_run(f'{idx}. ')
            num_run.bold = True
            num_run.font.size = Pt(11)
            
            text_run = p.add_run(text)
            text_run.font.color.rgb = RGBColor(0, 0, 255)
            text_run.font.size = Pt(11)
        
        doc.add_paragraph()
    else:
        doc.add_paragraph('[Aucun texte detecte]')
        doc.add_paragraph()

# Sauvegarder
output = folder / 'Chapitre_35_COMPLET.docx'
doc.save(str(output))

elapsed = time.time() - start_time
minutes = int(elapsed // 60)
seconds = int(elapsed % 60)

print('\n' + '='*70)
print(f'EXTRACTION TERMINEE !'.center(70))
print('='*70)
print(f'Total: {total_bubbles} bulles extraites de {len(images)} pages')
print(f'Temps: {minutes}m {seconds}s')
print(f'Fichier: {output}')
print('='*70)
