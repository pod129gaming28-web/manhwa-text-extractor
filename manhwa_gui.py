#!/usr/bin/env python3
"""
Manhwa Text Extractor - Interface Graphique
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
from pathlib import Path
import cv2
import numpy as np
from paddleocr import PaddleOCR
from docx import Document
from docx.shared import Pt, RGBColor
import os
import re


class ManhwaExtractorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("📚 Manhwa Text Extractor")
        self.root.geometry("700x600")
        self.root.resizable(False, False)
        
        self.folder_path = tk.StringVar()
        self.status_text = tk.StringVar(value="Prêt")
        self.progress_value = tk.DoubleVar(value=0)
        self.korean_enabled = tk.BooleanVar(value=True)
        self.english_enabled = tk.BooleanVar(value=False)
        self.is_processing = False
        
        self.ocr_ko = None
        self.ocr_en = None
        
        self.setup_ui()
    
    def setup_ui(self):
        style = ttk.Style()
        style.theme_use('clam')
        
        # Header
        header_frame = tk.Frame(self.root, bg="#2C3E50", height=80)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        title_label = tk.Label(
            header_frame,
            text="📚 Manhwa Text Extractor",
            font=("Segoe UI", 20, "bold"),
            bg="#2C3E50",
            fg="white"
        )
        title_label.pack(pady=25)
        
        # Main
        main_frame = tk.Frame(self.root, bg="#ECF0F1")
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Dossier
        folder_frame = tk.LabelFrame(
            main_frame,
            text="📁 Dossier à traiter",
            font=("Segoe UI", 11, "bold"),
            bg="#ECF0F1",
            fg="#2C3E50"
        )
        folder_frame.pack(fill=tk.X, pady=(0, 15))
        
        folder_entry = tk.Entry(
            folder_frame,
            textvariable=self.folder_path,
            font=("Segoe UI", 10),
            state="readonly",
            bg="white"
        )
        folder_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=10, pady=10)
        
        browse_btn = tk.Button(
            folder_frame,
            text="Parcourir...",
            command=self.browse_folder,
            bg="#3498DB",
            fg="white",
            font=("Segoe UI", 10, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            padx=20,
            pady=5
        )
        browse_btn.pack(side=tk.RIGHT, padx=10, pady=10)
        
        # Zone drop
        drop_frame = tk.Frame(main_frame, bg="white", relief=tk.RIDGE, borderwidth=2)
        drop_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        drop_label = tk.Label(
            drop_frame,
            text="📂\n\nGlissez un dossier ici\nou cliquez sur 'Parcourir'",
            font=("Segoe UI", 12),
            bg="white",
            fg="#7F8C8D",
            cursor="hand2"
        )
        drop_label.pack(expand=True)
        drop_label.bind("<Button-1>", lambda e: self.browse_folder())
        
        # Options
        options_frame = tk.LabelFrame(
            main_frame,
            text="⚙️ Options",
            font=("Segoe UI", 11, "bold"),
            bg="#ECF0F1",
            fg="#2C3E50"
        )
        options_frame.pack(fill=tk.X, pady=(0, 15))
        
        options_inner = tk.Frame(options_frame, bg="#ECF0F1")
        options_inner.pack(padx=10, pady=10)
        
        korean_check = tk.Checkbutton(
            options_inner,
            text="🇰🇷 Coréen",
            variable=self.korean_enabled,
            font=("Segoe UI", 10),
            bg="#ECF0F1"
        )
        korean_check.pack(side=tk.LEFT, padx=20)
        
        english_check = tk.Checkbutton(
            options_inner,
            text="🇬🇧 Anglais",
            variable=self.english_enabled,
            font=("Segoe UI", 10),
            bg="#ECF0F1"
        )
        english_check.pack(side=tk.LEFT, padx=20)
        
        # Bouton extraction
        extract_btn = tk.Button(
            main_frame,
            text="▶ EXTRAIRE LE TEXTE",
            command=self.start_extraction,
            bg="#27AE60",
            fg="white",
            font=("Segoe UI", 14, "bold"),
            cursor="hand2",
            relief=tk.FLAT,
            pady=15
        )
        extract_btn.pack(fill=tk.X, pady=(0, 15))
        
        # Progression
        progress_frame = tk.LabelFrame(
            main_frame,
            text="📊 Progression",
            font=("Segoe UI", 11, "bold"),
            bg="#ECF0F1",
            fg="#2C3E50"
        )
        progress_frame.pack(fill=tk.X, pady=(0, 15))
        
        self.progress_bar = ttk.Progressbar(
            progress_frame,
            variable=self.progress_value,
            maximum=100,
            mode='determinate'
        )
        self.progress_bar.pack(fill=tk.X, padx=10, pady=10)
        
        # Status
        status_label = tk.Label(
            main_frame,
            textvariable=self.status_text,
            font=("Segoe UI", 10),
            bg="#ECF0F1",
            fg="#2C3E50",
            anchor="w"
        )
        status_label.pack(fill=tk.X)
        
        # Footer
        footer = tk.Label(
            self.root,
            text="Manhwa Text Extractor v1.0",
            font=("Segoe UI", 8),
            bg="#ECF0F1",
            fg="#7F8C8D"
        )
        footer.pack(side=tk.BOTTOM, pady=5)
    
    def browse_folder(self):
        folder = filedialog.askdirectory(
            title="Sélectionnez le dossier du manhwa"
        )
        if folder:
            self.folder_path.set(folder)
            path = Path(folder)
            images = list(path.glob('*.jpeg')) + list(path.glob('*.jpg')) + list(path.glob('*.png'))
            self.status_text.set(f"✓ {len(images)} images trouvées")
    
    def start_extraction(self):
        if self.is_processing:
            messagebox.showwarning("En cours", "Extraction déjà en cours !")
            return
        
        if not self.folder_path.get():
            messagebox.showerror("Erreur", "Sélectionnez un dossier !")
            return
        
        if not self.korean_enabled.get() and not self.english_enabled.get():
            messagebox.showerror("Erreur", "Sélectionnez au moins une langue !")
            return
        
        self.is_processing = True
        self.progress_value.set(0)
        
        thread = threading.Thread(target=self.extract_text, daemon=True)
        thread.start()
    
    def extract_text(self):
        try:
            self.status_text.set("⏳ Initialisation...")
            
            if self.korean_enabled.get() and self.ocr_ko is None:
                self.ocr_ko = PaddleOCR(
                    lang='korean',
                    use_textline_orientation=True,
                    text_det_thresh=0.25,
                    text_det_box_thresh=0.55,
                    text_det_unclip_ratio=1.6
                )
            
            if self.english_enabled.get() and self.ocr_en is None:
                self.ocr_en = PaddleOCR(
                    lang='en',
                    use_textline_orientation=True,
                    text_det_thresh=0.25,
                    text_det_box_thresh=0.55,
                    text_det_unclip_ratio=1.6
                )
            
            folder = Path(self.folder_path.get())
            images = sorted(list(folder.glob('*.jpeg')) + list(folder.glob('*.jpg')) + list(folder.glob('*.png')))
            images = sorted(set(images))
            
            if not images:
                self.status_text.set("❌ Aucune image")
                self.is_processing = False
                return
            
            doc = Document()
            doc.add_heading('Extraction Manhwa', 0)
            doc.add_paragraph(f'{len(images)} pages')
            doc.add_paragraph()
            
            total_bubbles = 0
            
            for idx, img_path in enumerate(images, 1):
                self.status_text.set(f"📄 Page {idx}/{len(images)}")
                self.progress_value.set((idx / len(images)) * 100)
                
                bubbles = self.process_image(img_path)
                total_bubbles += len(bubbles)
                
                doc.add_heading(f'Page {idx}', level=2)
                
                if bubbles:
                    for num, text in enumerate(bubbles, 1):
                        p = doc.add_paragraph()
                        num_run = p.add_run(f'{num}. ')
                        num_run.bold = True
                        num_run.font.size = Pt(11)
                        
                        text_run = p.add_run(text)
                        text_run.font.color.rgb = RGBColor(0, 0, 255)
                        text_run.font.size = Pt(11)
                    
                    doc.add_paragraph()
                else:
                    doc.add_paragraph('[Aucun texte]')
                    doc.add_paragraph()
            
            output = folder / f'{folder.name}_extraction.docx'
            doc.save(str(output))
            
            self.progress_value.set(100)
            self.status_text.set(f"✓ Terminé ! {total_bubbles} bulles")
            
            messagebox.showinfo(
                "Terminé !",
                f"✓ {total_bubbles} bulles extraites\n\nFichier :\n{output}"
            )
            
        except Exception as e:
            self.status_text.set(f"❌ Erreur")
            messagebox.showerror("Erreur", str(e))
        
        finally:
            self.is_processing = False
    
    def process_image(self, img_path):
        img = cv2.imread(str(img_path))
        if img is None:
            return []
        
        h, w = img.shape[:2]
        
        chunks = []
        chunk_height = 3000
        overlap = 200
        
        for y in range(0, h, chunk_height - overlap):
            y_end = min(y + chunk_height, h)
            chunk = img[y:y_end, :]
            
            temp_path = f'temp_chunk_{len(chunks)}.jpg'
            cv2.imwrite(temp_path, chunk)
            
            preprocessed = self.light_preprocess(chunk)
            if preprocessed is not None:
                temp_prep = f'temp_prep_{len(chunks)}.jpg'
                cv2.imwrite(temp_prep, preprocessed)
                chunks.append((temp_prep, temp_path))
            else:
                chunks.append((temp_path, temp_path))
        
        all_texts = []
        
        for prep_path, orig_path in chunks:
            try:
                if self.korean_enabled.get():
                    result = self.ocr_ko.predict(prep_path)
                    texts = self.extract_texts(result)
                    all_texts.extend(texts)
                
                if self.english_enabled.get():
                    result = self.ocr_en.predict(prep_path)
                    texts = self.extract_texts(result)
                    all_texts.extend(texts)
                
            finally:
                if os.path.exists(prep_path):
                    os.remove(prep_path)
                if os.path.exists(orig_path) and orig_path != prep_path:
                    os.remove(orig_path)
        
        return all_texts
    
    def light_preprocess(self, img):
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        clahe = cv2.createCLAHE(clipLimit=1.5, tileGridSize=(8,8))
        enhanced = clahe.apply(gray)
        denoised = cv2.fastNlMeansDenoising(enhanced, h=7)
        return denoised
    
    def extract_texts(self, result):
        texts = []
        
        if result and len(result) > 0:
            for page_result in result:
                if 'rec_texts' in page_result:
                    for text, score in zip(page_result['rec_texts'], page_result['rec_scores']):
                        if score > 0.70 and text.strip() and not self.is_noise(text):
                            texts.append(text.strip())
        
        return texts
    
    def is_noise(self, text):
        text = text.strip()
        if len(text) < 2:
            return True
        
        if re.match(r'^[0-9<>|\.\-_~*:\'"`]+$', text):
            return True
        
        if text in ['UU', 'Mdo', '00000', '|0', '::', 'H']:
            return True
        
        return False


def main():
    root = tk.Tk()
    app = ManhwaExtractorGUI(root)
    root.mainloop()


if __name__ == '__main__':
    main()